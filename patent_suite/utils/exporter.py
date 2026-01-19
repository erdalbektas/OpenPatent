import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT

def export_patent_application(session_id, drafts_dir, final_export_dir):
    """
    Combines text files in /drafts into a standard .docx patent application.
    Includes standard headers and simulated line numbers.
    """
    print(f"Exporting professional .docx patent application for {session_id}...")
    
    if not os.path.exists(final_export_dir):
        os.makedirs(final_export_dir, exist_ok=True)
        
    doc = Document()
    
    # --- USPTO Standard Formatting (Phase 11) ---
    section = doc.sections[0]
    section.page_height = Cm(27.94) # Letter: 11 inches
    section.page_width = Cm(21.59)  # Letter: 8.5 inches
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set default style (Times New Roman 12pt)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5 # Improved readability

    # Title
    title = "LASER-BASED PRECISION TOASTING SYSTEM AND METHOD" # Default title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sections = [
        ("FIELD OF THE INVENTION", "field_of_the_invention.txt"),
        ("BACKGROUND OF THE INVENTION", "background_of_the_invention.txt"),
        ("SUMMARY OF THE INVENTION", "summary_of_the_invention.txt"),
        ("DETAILED DESCRIPTION", "detailed_description.txt"),
        ("CLAIMS", "claims.txt")
    ]
    
    line_count = 1
    
    for header, filename in sections:
        # Add Header
        doc.add_heading(header, level=1)
        
        content = ""
        file_path = os.path.join(drafts_dir, filename)
        if os.path.exists(file_path):
            with open(file_path, 'r') as f:
                content = f.read()
        else:
            content = f"[Drafting for {header} is currently in progress.]\n"

        # Add content with strict line numbering (every 5 lines)
        lines = content.split('\n')
        for line in lines:
            if line.strip() == "":
                doc.add_paragraph("")
                continue
                
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(-0.5) # Pull numbers into the margin slighty
            
            # Line numbering logic (Task 6.1)
            if line_count % 5 == 0:
                run = p.add_run(f"{line_count:<4}") # Left-aligned in margin area
                run.font.name = 'Courier New' # Distinctive technical look
                run.font.size = Pt(8)
                run.font.color.rgb = None # Standard black
            else:
                p.add_run("    ") # Alignment padding
                
            p.add_run(line)
            line_count += 1
            
    final_path = os.path.join(final_export_dir, f"Patent_Application_{session_id}.docx")
    doc.save(final_path)
    
    print(f"Professional .docx application exported to: {final_path}")
    return final_path

if __name__ == "__main__":
    # Mock run
    sample_drafts = "workspaces/test_v1/drafts"
    sample_export = "workspaces/test_v1/final_export"
    export_patent_application("test_v1", sample_drafts, sample_export)
